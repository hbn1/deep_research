#!/usr/bin/env python3
"""Enterprise-grade checks for runtime RAG document ingestion."""

from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from fastapi import FastAPI
from fastapi.testclient import TestClient


sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "app"))


class FakeRag:
    def __init__(self, chunks: int = 3):
        self.chunks = chunks
        self.ingested: list[dict] = []
        self.searched: list[dict] = []

    def ingest_bytes(self, content, filename, metadata=None):
        self.ingested.append({"content": content, "filename": filename, "metadata": metadata or {}})
        return self.chunks

    def get_collection_stats(self):
        return {"exists": True, "collection": "test_collection", "num_entities": 7}

    def search_records(self, query, k=5):
        self.searched.append({"query": query, "k": k})
        return [{"title": "Upload Guide", "doc_id": "rag://tenant/doc/file.pdf", "snippet": "answer", "score": 0.9}]


def fake_config(**overrides):
    values = {
        "api_key": "test-key",
        "tenant_id": "default_tenant",
        "user_id": "default_user",
        "milvus_host": "127.0.0.1",
        "milvus_port": 19530,
        "milvus_collection": "test_collection",
        "enable_milvus": True,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


class RagDocumentServiceTests(unittest.TestCase):
    def _service(self, tmp: str, max_upload_bytes: int = 1024):
        from backend.service.rag_document_service import RagDocumentService

        return RagDocumentService(
            config_path=str(Path(tmp) / "config.json"),
            upload_dir=Path(tmp) / "uploads",
            max_upload_bytes=max_upload_bytes,
            allowed_extensions={".pdf", ".docx"},
        )

    def test_service_success_persists_file_manifest_and_metadata(self):
        fake_rag = FakeRag(chunks=4)
        with tempfile.TemporaryDirectory() as tmp:
            service = self._service(tmp)
            with (
                patch("backend.service.rag_document_service.AppConfig.from_file", return_value=fake_config()),
                patch("backend.service.rag_document_service.configure_dashscope_endpoint") as endpoint,
                patch("backend.service.rag_document_service.init_rag_system") as init_rag,
                patch("backend.service.rag_document_service.RAGManager.get", return_value=fake_rag),
            ):
                record = service.ingest_document(
                    filename="../../知识 库.pdf",
                    content_type="application/pdf",
                    content=b"%PDF fake selectable text",
                    tenant_id="tenant / A",
                    user_id="user / A",
                    thread_id="thread-1",
                )

            self.assertEqual(record.filename, "知识 库.pdf")
            self.assertEqual(record.tenant_id, "tenant_A")
            self.assertEqual(record.user_id, "user_A")
            self.assertEqual(record.chunks, 4)
            self.assertTrue(Path(record.stored_path).is_file())
            self.assertTrue(str(Path(record.stored_path).resolve()).startswith(str((Path(tmp) / "uploads").resolve())))
            self.assertEqual(Path(record.stored_path).read_bytes(), b"%PDF fake selectable text")
            self.assertIn("rag://tenant_A/", record.source)
            self.assertEqual(fake_rag.ingested[0]["filename"], record.source)
            metadata = fake_rag.ingested[0]["metadata"]
            self.assertEqual(metadata["filename"], "知识 库.pdf")
            self.assertEqual(metadata["tenant_id"], "tenant_A")
            self.assertEqual(metadata["user_id"], "user_A")
            self.assertEqual(metadata["thread_id"], "thread-1")
            self.assertEqual(metadata["source"], record.source)
            endpoint.assert_called_once()
            init_rag.assert_called_once()

            manifest = Path(tmp) / "uploads" / "manifest.jsonl"
            lines = manifest.read_text(encoding="utf-8").splitlines()
            self.assertEqual(len(lines), 1)
            self.assertEqual(json.loads(lines[0])["doc_id"], record.doc_id)

    def test_service_rejects_unsupported_empty_and_oversized_files_before_config_load(self):
        from backend.service.rag_document_service import RagDocumentError

        with tempfile.TemporaryDirectory() as tmp:
            service = self._service(tmp, max_upload_bytes=3)
            with patch("backend.service.rag_document_service.AppConfig.from_file") as config_loader:
                with self.assertRaises(RagDocumentError) as unsupported:
                    service.ingest_document(
                        filename="notes.exe",
                        content_type="application/octet-stream",
                        content=b"123",
                        tenant_id="tenant",
                        user_id="user",
                    )
                self.assertEqual(unsupported.exception.status_code, 415)

                with self.assertRaises(RagDocumentError) as empty:
                    service.ingest_document(
                        filename="empty.pdf",
                        content_type="application/pdf",
                        content=b"",
                        tenant_id="tenant",
                        user_id="user",
                    )
                self.assertEqual(empty.exception.status_code, 400)

                with self.assertRaises(RagDocumentError) as too_large:
                    service.ingest_document(
                        filename="large.pdf",
                        content_type="application/pdf",
                        content=b"1234",
                        tenant_id="tenant",
                        user_id="user",
                    )
                self.assertEqual(too_large.exception.status_code, 413)
                config_loader.assert_not_called()

    def test_service_maps_rag_failures_and_empty_extraction_to_http_ready_errors(self):
        from backend.service.rag_document_service import RagDocumentError

        with tempfile.TemporaryDirectory() as tmp:
            service = self._service(tmp)
            with (
                patch("backend.service.rag_document_service.AppConfig.from_file", return_value=fake_config()),
                patch("backend.service.rag_document_service.configure_dashscope_endpoint"),
                patch("backend.service.rag_document_service.init_rag_system"),
                patch("backend.service.rag_document_service.RAGManager.get", side_effect=RuntimeError("milvus down")),
            ):
                with self.assertRaises(RagDocumentError) as failed:
                    service.ingest_document(
                        filename="file.pdf",
                        content_type="application/pdf",
                        content=b"%PDF abc",
                        tenant_id="tenant",
                        user_id="user",
                    )
                self.assertEqual(failed.exception.status_code, 502)

            with (
                patch("backend.service.rag_document_service.AppConfig.from_file", return_value=fake_config()),
                patch("backend.service.rag_document_service.configure_dashscope_endpoint"),
                patch("backend.service.rag_document_service.init_rag_system"),
                patch("backend.service.rag_document_service.RAGManager.get", return_value=FakeRag(chunks=0)),
            ):
                with self.assertRaises(RagDocumentError) as empty_chunks:
                    service.ingest_document(
                        filename="scan.pdf",
                        content_type="application/pdf",
                        content=b"%PDF abc",
                        tenant_id="tenant",
                        user_id="user",
                    )
                self.assertEqual(empty_chunks.exception.status_code, 422)

    def test_service_rejects_type_spoofing_and_tenant_quota(self):
        from backend.service.rag_document_service import RagDocumentError

        with tempfile.TemporaryDirectory() as tmp:
            service = self._service(tmp, max_upload_bytes=100)
            with self.assertRaises(RagDocumentError) as spoofed:
                service.ingest_document(
                    filename="fake.pdf",
                    content_type="application/pdf",
                    content=b"not a pdf",
                    tenant_id="tenant",
                    user_id="user",
                )
            self.assertEqual(spoofed.exception.status_code, 415)

            quota_service = self._service(tmp, max_upload_bytes=100)
            quota_service._max_tenant_storage_bytes = 8
            upload_dir = Path(tmp) / "uploads"
            upload_dir.mkdir(parents=True, exist_ok=True)
            manifest = upload_dir / "manifest.jsonl"
            manifest.write_text(
                json.dumps(self._record_dict("existing", "tenant"), ensure_ascii=False) + "\n",
                encoding="utf-8",
            )
            with (
                patch("backend.service.rag_document_service.AppConfig.from_file", return_value=fake_config()),
                patch("backend.service.rag_document_service.configure_dashscope_endpoint"),
            ):
                with self.assertRaises(RagDocumentError) as quota:
                    quota_service.ingest_document(
                        filename="quota.pdf",
                        content_type="application/pdf",
                        content=b"%PDF quota",
                        tenant_id="tenant",
                        user_id="user",
                    )
            self.assertEqual(quota.exception.status_code, 413)

    def test_document_listing_filters_tenant_skips_bad_manifest_and_honors_limit(self):
        with tempfile.TemporaryDirectory() as tmp:
            service = self._service(tmp)
            upload_dir = Path(tmp) / "uploads"
            upload_dir.mkdir(parents=True)
            manifest = upload_dir / "manifest.jsonl"
            rows = [
                self._record_dict("a1", "tenant_a"),
                {"broken": True},
                self._record_dict("b1", "tenant_b"),
                self._record_dict("a2", "tenant_a"),
            ]
            manifest.write_text("\n".join(json.dumps(row) for row in rows), encoding="utf-8")

            docs = service.list_documents("tenant_a", limit=1)
            self.assertEqual(len(docs), 1)
            self.assertEqual(docs[0].doc_id, "a2")

    def test_status_and_search_delegate_to_runtime_rag_instance(self):
        fake_rag = FakeRag()
        with tempfile.TemporaryDirectory() as tmp:
            service = self._service(tmp)
            with (
                patch("backend.service.rag_document_service.AppConfig.from_file", return_value=fake_config()),
                patch("backend.service.rag_document_service.configure_dashscope_endpoint"),
                patch("backend.service.rag_document_service.RAGManager.get_or_none", return_value=fake_rag),
            ):
                status = service.status("tenant / A")
            self.assertEqual(status["tenant_id"], "tenant_A")
            self.assertTrue(status["runtime_initialized"])
            self.assertEqual(status["stats"]["num_entities"], 7)

            with (
                patch("backend.service.rag_document_service.AppConfig.from_file", return_value=fake_config()),
                patch("backend.service.rag_document_service.configure_dashscope_endpoint"),
                patch("backend.service.rag_document_service.RAGManager.get", return_value=fake_rag),
            ):
                results = service.search("how to upload?", tenant_id="tenant / A", limit=2)
            self.assertEqual(results[0]["snippet"], "answer")
            self.assertEqual(fake_rag.searched[-1], {"query": "how to upload?", "k": 2})

    @staticmethod
    def _record_dict(doc_id: str, tenant_id: str):
        return {
            "doc_id": doc_id,
            "filename": f"{doc_id}.pdf",
            "source": f"rag://{tenant_id}/{doc_id}/{doc_id}.pdf",
            "tenant_id": tenant_id,
            "user_id": "user",
            "content_type": "application/pdf",
            "size_bytes": 10,
            "chunks": 1,
            "collection": "test_collection",
            "stored_path": f"uploads/{doc_id}.pdf",
            "uploaded_at": "2026-07-07T00:00:00",
        }


class RagRouterApiTests(unittest.TestCase):
    def _client(self, fake_service):
        from backend.config import AppSettings
        from backend.dependencies.admin import get_runtime_settings
        from backend.router.rag_router import router
        from backend.service import get_rag_document_service

        app = FastAPI()
        app.include_router(router)
        app.dependency_overrides[get_rag_document_service] = lambda: fake_service
        app.dependency_overrides[get_runtime_settings] = lambda: AppSettings(admin_api_required=False)
        return TestClient(app)

    def test_upload_route_accepts_multipart_file(self):
        from backend.schemas.rag import RagDocumentRecord

        class FakeService:
            def ingest_document(self, **kwargs):
                self.kwargs = kwargs
                return RagDocumentRecord(**RagDocumentServiceTests._record_dict("doc_test", kwargs["tenant_id"]))

        fake_service = FakeService()
        response = self._client(fake_service).post(
            "/api/v1/rag/documents",
            data={"tenant_id": "tenant_a", "user_id": "user_a", "thread_id": "thread_a"},
            files={
                "file": (
                    "example.docx",
                    b"fake docx bytes",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body["status"], "indexed")
        self.assertEqual(body["document"]["doc_id"], "doc_test")
        self.assertEqual(fake_service.kwargs["tenant_id"], "tenant_a")
        self.assertEqual(fake_service.kwargs["thread_id"], "thread_a")

    def test_router_maps_service_errors_and_validates_search_payload(self):
        from backend.service.rag_document_service import RagDocumentError

        class FakeService:
            def ingest_document(self, **kwargs):
                raise RagDocumentError("too large", 413)

            def search(self, **kwargs):
                return []

        client = self._client(FakeService())
        upload = client.post(
            "/api/v1/rag/documents",
            files={"file": ("large.pdf", b"123", "application/pdf")},
        )
        self.assertEqual(upload.status_code, 413)
        self.assertIn("too large", upload.json()["detail"])

        invalid = client.post("/api/v1/rag/search", json={"query": "", "limit": 5})
        self.assertEqual(invalid.status_code, 422)

        invalid_limit = client.post("/api/v1/rag/search", json={"query": "x", "limit": 50})
        self.assertEqual(invalid_limit.status_code, 422)

    def test_router_lists_status_and_searches_documents(self):
        class FakeService:
            def list_documents(self, tenant_id="", limit=50):
                return []

            def status(self, tenant_id="default_tenant"):
                return {
                    "tenant_id": tenant_id,
                    "collection": "test_collection",
                    "milvus_host": "127.0.0.1",
                    "milvus_port": 19530,
                    "configured_enabled": True,
                    "runtime_initialized": True,
                    "stats": {"num_entities": 2},
                    "documents": [],
                }

            def search(self, query, tenant_id="default_tenant", limit=5):
                return [{"snippet": query, "score": 1.0}]

        client = self._client(FakeService())
        self.assertEqual(client.get("/api/v1/rag/documents?tenant_id=t1").json()["documents"], [])
        status = client.get("/api/v1/rag/status?tenant_id=t1").json()
        self.assertEqual(status["tenant_id"], "t1")
        search = client.post("/api/v1/rag/search", json={"query": "needle", "tenant_id": "t1", "limit": 1})
        self.assertEqual(search.status_code, 200)
        self.assertEqual(search.json()["results"][0]["snippet"], "needle")


class RagCoreExtractionTests(unittest.TestCase):
    def test_collection_checks_use_rag_connection_alias(self):
        from app.mult_agents.rag.core import RAGConfig, RAGSystem

        rag = object.__new__(RAGSystem)
        rag.config = RAGConfig(collection_name="alias_collection")
        rag._connection_alias = "rag_alias_for_test"

        with patch("app.mult_agents.rag.core.utility.has_collection", return_value=False) as has_collection:
            self.assertEqual(rag.search_records("anything", k=1), [])
            has_collection.assert_called_with("alias_collection", using="rag_alias_for_test")

        with patch("app.mult_agents.rag.core.utility.has_collection", return_value=False) as has_collection:
            status = rag.get_collection_stats()
            self.assertFalse(status["exists"])
            has_collection.assert_called_with("alias_collection", using="rag_alias_for_test")

    def test_docx_extraction_includes_paragraphs_and_table_cells(self):
        import docx
        from app.mult_agents.rag.core import RAGSystem

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.docx"
            doc = docx.Document()
            doc.add_paragraph("Enterprise upload guide")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Owner"
            table.cell(0, 1).text = "RAG Team"
            doc.save(path)

            text = RAGSystem._extract_docx_text(path.read_bytes())
            self.assertIn("Enterprise upload guide", text)
            self.assertIn("Owner | RAG Team", text)

    def test_ingest_text_adds_chunk_metadata_without_real_milvus(self):
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        from app.mult_agents.rag.core import RAGSystem

        rag = object.__new__(RAGSystem)
        rag.md_splitter = None
        rag.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=20,
            chunk_overlap=0,
            length_function=len,
            separators=[" ", ""],
        )
        captured = {}

        def fake_add_documents(docs):
            captured["docs"] = docs
            return len(docs)

        rag.add_documents = fake_add_documents
        chunks = rag.ingest_text(
            "alpha beta gamma delta epsilon zeta eta theta",
            source="unit.txt",
            metadata={"tenant_id": "tenant_a"},
        )

        docs = captured["docs"]
        self.assertGreater(chunks, 1)
        self.assertEqual(chunks, len(docs))
        self.assertEqual([doc.metadata["chunk_index"] for doc in docs], list(range(len(docs))))
        self.assertTrue(all(doc.metadata["total_chunks"] == len(docs) for doc in docs))
        self.assertTrue(all(doc.metadata["tenant_id"] == "tenant_a" for doc in docs))


if __name__ == "__main__":
    unittest.main()
