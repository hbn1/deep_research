"""Enterprise RAG Core Module."""
from __future__ import annotations

import logging
import threading
import hashlib
import time
import re
import math as _math
from collections import OrderedDict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Optional

from langchain_community.embeddings import DashScopeEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from pymilvus import connections, utility

logger = logging.getLogger(__name__)

try:
    from langchain_milvus import Milvus as _MilvusVectorStore
    _MILVUS_BACKEND = "langchain_milvus"
except ImportError:
    from langchain_community.vectorstores import Milvus as _MilvusVectorStore
    _MILVUS_BACKEND = "langchain_community"

_HAS_BM25 = False
try:
    from pymilvus.model.sparse import BM25EmbeddingFunction
    _HAS_BM25 = True
except ImportError:
    pass
class _SimpleBM25Scorer:
    """Pure Python BM25 scorer for keyword retrieval."""
    def __init__(self, k1=1.5, b=0.75):
        self.k1 = k1
        self.b = b
        self._corpus = []
        self._doc_count = 0
        self._avgdl = 0.0
        self._idf = {}
        self._doc_freqs = {}

    def _tokenize(self, text):
        tokens = []
        for token in re.findall(r"[a-zA-Z0-9_]+|\S", text):
            token = token.strip()
            if token:
                tokens.append(token.lower())
        return tokens

    def index(self, documents):
        self._corpus = [{"tokens": self._tokenize(d), "text": d} for d in documents]
        self._doc_count = len(documents)
        total_len = sum(len(d["tokens"]) for d in self._corpus)
        self._avgdl = total_len / max(self._doc_count, 1)
        self._doc_freqs.clear()
        for doc_entry in self._corpus:
            seen = set()
            for token in doc_entry["tokens"]:
                if token not in seen:
                    self._doc_freqs[token] = self._doc_freqs.get(token, 0) + 1
                    seen.add(token)
        self._idf.clear()
        for token, df in self._doc_freqs.items():
            self._idf[token] = _math.log((self._doc_count - df + 0.5) / (df + 0.5) + 1.0)

    def search(self, query, k=10):
        query_tokens = self._tokenize(query)
        scores = []
        for idx, doc_entry in enumerate(self._corpus):
            doc_tokens = doc_entry["tokens"]
            doc_len = len(doc_tokens)
            score = 0.0
            tf_map = {}
            for t in doc_tokens:
                tf_map[t] = tf_map.get(t, 0) + 1
            for token in query_tokens:
                if token not in self._idf:
                    continue
                tf = tf_map.get(token, 0)
                idf = self._idf[token]
                numerator = tf * (self.k1 + 1)
                denominator = tf + self.k1 * (1 - self.b + self.b * doc_len / max(self._avgdl, 1))
                score += idf * numerator / max(denominator, 1e-8)
            scores.append((idx, score))
        scores.sort(key=lambda x: x[1], reverse=True)
        return scores[:k]


@dataclass(frozen=True)
class RAGConfig:
    """RAG system configuration (frozen/immutable)."""
    milvus_host: str = "127.0.0.1"
    milvus_port: int = 19530
    collection_name: str = "mult_agent_knowledge"
    embedding_model: str = "text-embedding-v1"
    chunk_size: int = 800
    chunk_overlap: int = 150
    use_semantic_chunking: bool = True
    chunk_context_expand: int = 1
    retrieval_top_k: int = 10
    final_top_k: int = 5
    use_hybrid_search: bool = True
    hybrid_dense_weight: float = 0.6
    similarity_threshold: float = 0.0
    enable_rerank: bool = True
    rerank_model: str = "gte-rerank"
    rerank_top_n: int = 5
    enable_query_rewrite: bool = True
    query_rewrite_max_queries: int = 2
    enable_cache: bool = True
    cache_ttl_seconds: int = 600
    cache_max_entries: int = 500


@dataclass
class RetrievalRecord:
    """Unified retrieval result record."""
    source_id: str = ""
    doc_id: str = ""
    title: str = ""
    snippet: str = ""
    source_type: str = "local"
    score: float = 0.0
    metadata: dict = field(default_factory=dict)
    chunk_index: int = 0
    total_chunks: int = 0

    def to_dict(self):
        return {
            "source_id": self.source_id, "doc_id": self.doc_id,
            "title": self.title, "snippet": self.snippet,
            "source_type": self.source_type, "score": self.score,
            "metadata": self.metadata,
        }


class _ThreadSafeLRUCache:
    """Thread-safe LRU cache for RAG queries."""
    def __init__(self, max_entries=500, ttl_seconds=600):
        self._max_entries = max_entries
        self._ttl_seconds = ttl_seconds
        self._cache = OrderedDict()
        self._lock = threading.Lock()

    def _make_key(self, query, k):
        return hashlib.md5(f"{query}::{k}".encode()).hexdigest()

    def get(self, query, k):
        key = self._make_key(query, k)
        with self._lock:
            if key not in self._cache:
                return None
            timestamp, value = self._cache[key]
            if time.time() - timestamp > self._ttl_seconds:
                del self._cache[key]
                return None
            self._cache.move_to_end(key)
            return value

    def set(self, query, k, records):
        key = self._make_key(query, k)
        with self._lock:
            if key in self._cache:
                self._cache.move_to_end(key)
            self._cache[key] = (time.time(), records)
            while len(self._cache) > self._max_entries:
                self._cache.popitem(last=False)

    def clear(self):
        with self._lock:
            self._cache.clear()


class RAGManager:
    """Thread-safe multi-tenant RAG instance manager."""
    _instances = {}
    _lock = threading.RLock()

    @classmethod
    def get(cls, api_key, config=None, tenant_id="default_tenant"):
        cfg = config or RAGConfig()
        key = f"{tenant_id}:{cfg.collection_name}:{cfg.milvus_host}:{cfg.milvus_port}"
        with cls._lock:
            if key not in cls._instances:
                cls._instances[key] = RAGSystem(api_key, cfg, instance_key=key)
            return cls._instances[key]

    @classmethod
    def get_or_none(cls, tenant_id="default_tenant", collection_name=""):
        with cls._lock:
            for key, instance in cls._instances.items():
                if tenant_id in key and collection_name in key:
                    return instance
        return None

    @classmethod
    def evict(cls, tenant_id="", collection_name=""):
        with cls._lock:
            keys_to_remove = [
                k for k in cls._instances
                if (not tenant_id or tenant_id in k)
                and (not collection_name or collection_name in k)
            ]
            for k in keys_to_remove:
                del cls._instances[k]
            logger.info("RAGManager evicted %d instance(s)", len(keys_to_remove))

    @classmethod
    def list_instances(cls):
        with cls._lock:
            return [
                {"key": k, "collection": v.config.collection_name,
                 "host": v.config.milvus_host, "port": v.config.milvus_port}
                for k, v in cls._instances.items()
            ]

class RAGSystem:
    """Enterprise RAG System with hybrid search, reranking, dynamic ingestion."""

    def __init__(self, api_key, config=None, instance_key="default"):
        self.config = config or RAGConfig()
        self.api_key = api_key
        self._instance_key = instance_key
        self._connection_alias = f"rag_{instance_key.replace(':', '_').replace('.', '_')}"

        self.embeddings = DashScopeEmbeddings(
            model=self.config.embedding_model,
            dashscope_api_key=self.api_key,
        )

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "?", "!", ",", " ", ""],
        )

        if self.config.use_semantic_chunking:
            try:
                self.md_splitter = MarkdownHeaderTextSplitter(
                    headers_to_split_on=[
                        ("#", "h1"), ("##", "h2"), ("###", "h3"), ("####", "h4"),
                    ],
                    strip_headers=False,
                )
            except Exception:
                self.md_splitter = None
        else:
            self.md_splitter = None

        self._connect_to_milvus()

        self.vectorstore = _MilvusVectorStore(
            embedding_function=self.embeddings,
            collection_name=self.config.collection_name,
            connection_args={
                "uri": f"http://{self.config.milvus_host}:{self.config.milvus_port}",
                "token": "",
            },
            auto_id=True,
        )

        self._bm25_scorer = None
        self._bm25_docs = []
        self._bm25_metadata = []
        self._bm25_needs_rebuild = True
        self._bm25_lock = threading.Lock()

        if self.config.enable_cache:
            self._cache = _ThreadSafeLRUCache(
                max_entries=self.config.cache_max_entries,
                ttl_seconds=self.config.cache_ttl_seconds,
            )
        else:
            self._cache = None

        logger.info(
            "RAG init | backend=%s | collection=%s | hybrid=%s | rerank=%s",
            _MILVUS_BACKEND, self.config.collection_name,
            self.config.use_hybrid_search, self.config.enable_rerank,
        )

    def _connect_to_milvus(self):
        try:
            try:
                if connections.has_connection(self._connection_alias):
                    connections.disconnect(self._connection_alias)
            except Exception:
                pass
            connections.connect(
                alias=self._connection_alias,
                host=self.config.milvus_host,
                port=self.config.milvus_port,
            )
        except Exception as exc:
            logger.error("Milvus connection failed: %s", exc)

    # === Public API ===

    def search(self, query, k=3):
        """Backward-compatible formatted string search."""
        try:
            records = self.search_records(query, k=k)
            if not records:
                return "No relevant information found."
            lines = ["Retrieved information:"]
            for idx, record in enumerate(records, 1):
                lines.append(f"{idx}. {record['snippet']}")
                lines.append(f"   (source: {record['doc_id']}, score: {record.get('score', 0):.3f})")
            return "\n".join(lines)
        except Exception as exc:
            logger.error("Search failed: %s", exc)
            return f"Search error: {str(exc)}"

    def search_records(
        self, query, k=5, metadata_filter=None,
        enable_hybrid=None, enable_rerank=None, enable_cache=None,
    ):
        """Full RAG pipeline: cache -> rewrite -> hybrid -> RRF -> rerank -> expand."""
        if not utility.has_collection(self.config.collection_name):
            logger.warning("Collection %s does not exist", self.config.collection_name)
            return []

        use_hybrid = enable_hybrid if enable_hybrid is not None else self.config.use_hybrid_search
        use_rerank = enable_rerank if enable_rerank is not None else self.config.enable_rerank
        use_cache = enable_cache if enable_cache is not None else self.config.enable_cache

        if use_cache and self._cache is not None:
            cached = self._cache.get(query, k)
            if cached is not None:
                return cached

        queries = self._rewrite_query(query) if self.config.enable_query_rewrite else [query]

        all_candidates = {}
        for q in queries:
            if use_hybrid:
                candidates = self._hybrid_search(q, k=self.config.retrieval_top_k, metadata_filter=metadata_filter)
            else:
                candidates = self._dense_search(q, k=self.config.retrieval_top_k, metadata_filter=metadata_filter)
            for rec in candidates:
                key = rec.doc_id or rec.snippet[:80]
                if key not in all_candidates or rec.score > all_candidates[key].score:
                    all_candidates[key] = rec

        if not all_candidates:
            return []

        if use_rerank and len(all_candidates) > 1:
            ranked = self._rerank(query, list(all_candidates.values()))
        else:
            ranked = sorted(all_candidates.values(), key=lambda x: x.score, reverse=True)

        top_k = ranked[:k]

        if self.config.chunk_context_expand > 0:
            top_k = self._expand_chunk_context(top_k)

        if self.config.similarity_threshold > 0:
            top_k = [r for r in top_k if r.score >= self.config.similarity_threshold]

        records = [r.to_dict() for r in top_k]

        if use_cache and self._cache is not None:
            self._cache.set(query, k, records)

        logger.info("RAG search | query=%.60s | hybrid=%s | rerank=%s | results=%d",
                     query, use_hybrid, use_rerank, len(records))
        return records

    # === Ingestion API ===

    def add_documents(self, documents):
        self.vectorstore.add_documents(documents)
        self._bm25_needs_rebuild = True
        logger.info("Added %d documents", len(documents))
        return len(documents)

    def ingest_text(self, text, source="", metadata=None, doc_type="text"):
        base_meta = {
            "source": source or "inline_text",
            "doc_type": doc_type,
            "ingested_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        }
        if metadata:
            base_meta.update(metadata)

        if doc_type == "markdown" and self.md_splitter is not None:
            docs = self._chunk_markdown(text, base_meta)
        else:
            docs = self.text_splitter.create_documents([text], metadatas=[base_meta])

        return self.add_documents(docs)

    def ingest_paths(self, paths):
        total = 0
        for path in paths:
            try:
                path = Path(path)
                suffix = path.suffix.lower()
                if suffix == ".md":
                    text = path.read_text(encoding="utf-8")
                    total += self.ingest_text(text, source=str(path), doc_type="markdown")
                elif suffix in (".txt", ".py", ".js", ".ts", ".java", ".go", ".rs"):
                    text = path.read_text(encoding="utf-8")
                    doc_type = "code" if suffix in (".py", ".js", ".ts", ".java", ".go", ".rs") else "text"
                    total += self.ingest_text(text, source=str(path), doc_type=doc_type)
                elif suffix == ".pdf":
                    total += self._ingest_pdf(path)
                elif suffix == ".docx":
                    total += self._ingest_docx(path)
                else:
                    logger.warning("Unsupported file type: %s", suffix)
            except Exception as exc:
                logger.error("Failed to ingest %s: %s", path, exc)
        return total

    def ingest_bytes(self, content, filename, metadata=None):
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return self._ingest_pdf_bytes(content, filename, metadata)
        elif suffix == ".docx":
            return self._ingest_docx_bytes(content, filename, metadata)
        else:
            text = content.decode("utf-8", errors="replace")
            doc_type = "markdown" if suffix == ".md" else "text"
            return self.ingest_text(text, source=filename, metadata=metadata, doc_type=doc_type)

    def delete_document(self, doc_id):
        try:
            expr = f'source == "{doc_id}"'
            self.vectorstore.delete(expr=expr)
            self._bm25_needs_rebuild = True
            logger.info("Deleted document: %s", doc_id)
            return True
        except Exception as exc:
            logger.error("Failed to delete %s: %s", doc_id, exc)
            return False

    def update_document(self, doc_id, new_text, metadata=None):
        self.delete_document(doc_id)
        return self.ingest_text(new_text, source=doc_id, metadata=metadata)

    def get_collection_stats(self):
        try:
            if not utility.has_collection(self.config.collection_name):
                return {"exists": False, "collection": self.config.collection_name}
            return {
                "exists": True, "collection": self.config.collection_name,
                "backend": _MILVUS_BACKEND,
                "host": self.config.milvus_host, "port": self.config.milvus_port,
            }
        except Exception as exc:
            return {"error": str(exc)}

    # === Internal: Retrieval Pipeline ===

    def _dense_search(self, query, k, metadata_filter=None):
        expr = self._build_filter_expr(metadata_filter) if metadata_filter else None
        try:
            if expr:
                docs_with_scores = self.vectorstore.similarity_search_with_score(query, k=k, expr=expr)
            else:
                docs_with_scores = self.vectorstore.similarity_search_with_score(query, k=k)
        except Exception:
            docs = self.vectorstore.similarity_search(query, k=k, expr=expr if expr else None)
            docs_with_scores = [(doc, 0.5) for doc in docs]
        return [self._doc_to_record(doc, score, idx) for idx, (doc, score) in enumerate(docs_with_scores)]

    def _hybrid_search(self, query, k, metadata_filter=None):
        dense_results = self._dense_search(query, k=max(k, 20), metadata_filter=metadata_filter)
        bm25_results = self._bm25_search(query, k=max(k, 20))
        if not bm25_results:
            return dense_results[:k]
        if not dense_results:
            return bm25_results[:k]
        return self._rrf_fusion(dense_results, bm25_results, dense_weight=self.config.hybrid_dense_weight, k=k)

    def _bm25_search(self, query, k):
        self._ensure_bm25_index()
        if self._bm25_scorer is None or not self._bm25_docs:
            return []
        try:
            scored = self._bm25_scorer.search(query, k=k)
        except Exception as exc:
            logger.warning("BM25 search failed: %s", exc)
            return []
        results = []
        for idx, (doc_idx, score) in enumerate(scored):
            if doc_idx >= len(self._bm25_docs):
                continue
            meta = self._bm25_metadata[doc_idx] if doc_idx < len(self._bm25_metadata) else {}
            results.append(RetrievalRecord(
                source_id=f"BM25-{idx}",
                doc_id=str(meta.get("source", "")),
                title=Path(str(meta.get("source", ""))).name or f"bm25_{idx}",
                snippet=self._bm25_docs[doc_idx][:self.config.chunk_size],
                source_type="local",
                score=score / max(1.0, score + 1.0),
                metadata=meta,
            ))
        return results

    def _rrf_fusion(self, dense, sparse, dense_weight=0.6, k=60):
        rrf_scores = {}
        record_map = {}
        sparse_weight = 1.0 - dense_weight

        for rank, rec in enumerate(dense, 1):
            key = rec.doc_id or rec.snippet[:80]
            rrf_scores[key] = rrf_scores.get(key, 0) + dense_weight / (k + rank)
            if key not in record_map or rec.score > record_map[key].score:
                record_map[key] = rec

        for rank, rec in enumerate(sparse, 1):
            key = rec.doc_id or rec.snippet[:80]
            rrf_scores[key] = rrf_scores.get(key, 0) + sparse_weight / (k + rank)
            if key not in record_map:
                record_map[key] = rec

        sorted_keys = sorted(rrf_scores, key=lambda x: rrf_scores[x], reverse=True)
        merged = []
        for key in sorted_keys:
            rec = record_map[key]
            rec.score = rrf_scores[key]
            merged.append(rec)
        return merged

    def _rerank(self, query, candidates):
        if not candidates:
            return candidates
        rerank_top_n = min(self.config.rerank_top_n, len(candidates))
        rerank_candidates = candidates[:rerank_top_n]
        try:
            reranked = self._dashscope_rerank(query, rerank_candidates)
            if reranked:
                reranked_ids = {id(r): r for r in reranked}
                return reranked + [c for c in candidates if id(c) not in reranked_ids]
        except Exception as exc:
            logger.warning("DashScope rerank failed: %s", exc)
        return self._overlap_rerank(query, candidates)

    def _dashscope_rerank(self, query, candidates):
        import urllib.request
        import json as _json
        documents = [c.snippet for c in candidates]
        payload = _json.dumps({
            "model": self.config.rerank_model,
            "input": {"query": query, "documents": documents},
            "parameters": {"top_n": len(candidates), "return_documents": False},
        }).encode("utf-8")
        req = urllib.request.Request(
            "https://dashscope.aliyuncs.com/api/v1/services/rerank/text-rerank/text-rerank",
            data=payload,
            headers={
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
            },
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = _json.loads(resp.read().decode())
        output = result.get("output", {})
        ranked_indices = output.get("results", [])
        if not ranked_indices:
            return []
        reranked = []
        for item in ranked_indices:
            idx = item.get("index", -1)
            relevance_score = item.get("relevance_score", 0.5)
            if 0 <= idx < len(candidates):
                rec = candidates[idx]
                rec.score = relevance_score
                reranked.append(rec)
        return reranked

    def _overlap_rerank(self, query, candidates):
        query_lower = query.lower()
        query_terms = set(re.findall(r"\w+", query_lower))
        for rec in candidates:
            snippet_lower = rec.snippet.lower()
            snippet_terms = set(re.findall(r"\w+", snippet_lower))
            if query_terms:
                overlap = len(query_terms & snippet_terms) / len(query_terms)
            else:
                overlap = 0.0
            rec.score = 0.4 * rec.score + 0.6 * overlap
        return sorted(candidates, key=lambda x: x.score, reverse=True)

    def _rewrite_query(self, query):
        variants = [query]
        max_v = self.config.query_rewrite_max_queries
        if "?" in query or "?" in query:
            dec = query.replace("?", "").replace("?", "").strip()
            if dec != query:
                variants.append(dec)
        stopwords = {"de", "shi", "le", "zai", "he", "ye", "dou", "jiu", "ma", "ne", "ba", "a"}
        keywords = " ".join([w for w in re.split(r"[\s,.,!,?,]+", query) if w and w.lower() not in stopwords])
        if keywords != query and len(keywords) > 2:
            variants.append(keywords)
        return variants[:max_v]

    def _expand_chunk_context(self, candidates):
        expand_n = self.config.chunk_context_expand
        if expand_n <= 0:
            return candidates
        for rec in candidates:
            chunk_idx = rec.metadata.get("chunk_index", -1)
            if chunk_idx < 0:
                continue
            source = rec.metadata.get("source", rec.doc_id)
            all_chunks = self._get_document_chunks(source)
            if not all_chunks:
                continue
            start = max(0, chunk_idx - expand_n)
            end = min(len(all_chunks), chunk_idx + expand_n + 1)
            expanded = " ".join(all_chunks[start:end])
            if len(expanded) > len(rec.snippet):
                rec.snippet = expanded
        return candidates

    def _get_document_chunks(self, source):
        try:
            expr = f'source == "{source}"'
            docs = self.vectorstore.similarity_search("", k=100, expr=expr)
            chunks_with_idx = []
            for doc in docs:
                meta = doc.metadata or {}
                idx = meta.get("chunk_index", 0)
                chunks_with_idx.append((idx, doc.page_content))
            chunks_with_idx.sort(key=lambda x: x[0])
            return [c[1] for c in chunks_with_idx]
        except Exception:
            return []

    def _build_filter_expr(self, metadata_filter):
        parts = []
        for key, value in metadata_filter.items():
            if isinstance(value, str):
                parts.append(f'{key} == "{value}"')
            elif isinstance(value, (int, float)):
                parts.append(f"{key} == {value}")
            elif isinstance(value, list):
                list_str = ", ".join(f'"{v}"' if isinstance(v, str) else str(v) for v in value)
                parts.append(f"{key} in [{list_str}]")
        return " && ".join(parts) if parts else ""

    # === Internal: Chunking & Ingestion ===

    def _chunk_markdown(self, text, base_meta):
        try:
            md_docs = self.md_splitter.split_text(text)
        except Exception:
            return self.text_splitter.create_documents([text], metadatas=[base_meta])

        final_docs = []
        for idx, md_doc in enumerate(md_docs):
            merged_meta = {**base_meta, **md_doc.metadata, "chunk_index": idx}
            if len(md_doc.page_content) > self.config.chunk_size * 1.5:
                sub_docs = self.text_splitter.create_documents(
                    [md_doc.page_content], metadatas=[merged_meta]
                )
                for sub_idx, sub_doc in enumerate(sub_docs):
                    sub_doc.metadata["chunk_index"] = idx * 1000 + sub_idx
                final_docs.extend(sub_docs)
            else:
                md_doc.metadata = merged_meta
                final_docs.append(md_doc)

        total = len(final_docs)
        for doc in final_docs:
            doc.metadata["total_chunks"] = total
        return final_docs

    def _ingest_pdf(self, path):
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            logger.warning("pypdf not installed: %s", path)
            return 0
        except Exception as exc:
            logger.error("PDF extraction failed: %s", exc)
            return 0
        return self.ingest_text(text, source=str(path), doc_type="text")

    def _ingest_docx(self, path):
        try:
            import docx
            doc = docx.Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text)
        except ImportError:
            logger.warning("python-docx not installed: %s", path)
            return 0
        except Exception as exc:
            logger.error("DOCX extraction failed: %s", exc)
            return 0
        return self.ingest_text(text, source=str(path), doc_type="text")

    def _ingest_pdf_bytes(self, content, filename, metadata=None):
        try:
            import pypdf, io
            reader = pypdf.PdfReader(io.BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            logger.warning("pypdf not installed")
            return 0
        except Exception as exc:
            logger.error("PDF bytes extraction failed: %s", exc)
            return 0
        return self.ingest_text(text, source=filename, metadata=metadata, doc_type="text")

    def _ingest_docx_bytes(self, content, filename, metadata=None):
        try:
            import docx, io
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs if p.text)
        except ImportError:
            logger.warning("python-docx not installed")
            return 0
        except Exception as exc:
            logger.error("DOCX bytes extraction failed: %s", exc)
            return 0
        return self.ingest_text(text, source=filename, metadata=metadata, doc_type="text")

    # === Internal: Utilities ===

    def _doc_to_record(self, doc, score, idx):
        metadata = doc.metadata or {}
        source = str(metadata.get("source") or "").strip()
        title = Path(source).name if source else f"local_chunk-{idx + 1}"
        return RetrievalRecord(
            source_id=f"LOC-{idx + 1}",
            doc_id=source,
            title=title,
            snippet=doc.page_content,
            source_type="local",
            score=float(score),
            metadata=metadata,
            chunk_index=int(metadata.get("chunk_index", idx)),
            total_chunks=int(metadata.get("total_chunks", 1)),
        )

    def _ensure_bm25_index(self):
        if not self._bm25_needs_rebuild:
            return
        with self._bm25_lock:
            if not self._bm25_needs_rebuild:
                return
            self._rebuild_bm25_index()
            self._bm25_needs_rebuild = False

    def _rebuild_bm25_index(self):
        try:
            docs = self.vectorstore.similarity_search("", k=10000)
        except Exception:
            docs = []
        if not docs:
            self._bm25_scorer = None
            self._bm25_docs = []
            self._bm25_metadata = []
            return
        self._bm25_docs = [d.page_content for d in docs]
        self._bm25_metadata = [d.metadata or {} for d in docs]
        self._bm25_scorer = _SimpleBM25Scorer()
        self._bm25_scorer.index(self._bm25_docs)
        logger.info("BM25 index rebuilt | docs=%d", len(self._bm25_docs))